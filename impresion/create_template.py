import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT

def generate_invoice_template(output_path):
    doc = Document()
    
    # Márgenes de la página (1 cm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Título Principal
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_title = p_header.add_run("FACTURA ELECTRÓNICA DE VENTA\n")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(29, 78, 216) # Azul corporativo

    run_num = p_header.add_run("No. {{ doc.num_documento }}\n")
    run_num.font.name = 'Arial'
    run_num.font.size = Pt(14)
    run_num.font.bold = True

    run_fch = p_header.add_run("Fecha: {{ doc.fch_documento }}\nResolución DIAN: {{ doc.resolucion }}")
    run_fch.font.name = 'Arial'
    run_fch.font.size = Pt(9)

    # Tabla de Datos de la Empresa y Cliente (2 Columnas)
    table_info = doc.add_table(rows=1, cols=2)
    table_info.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_info.autofit = False

    cell_empresa = table_info.rows[0].cells[0]
    cell_cliente = table_info.rows[0].cells[1]

    p_emp = cell_empresa.paragraphs[0]
    p_emp.add_run("EMISOR:\n").bold = True
    p_emp.add_run("{{ empresa.nom_empresa }}\nNIT: {{ empresa.nit }}\nDirección: {{ empresa.direccion }}")

    p_cli = cell_cliente.paragraphs[0]
    p_cli.add_run("ADQUIRENTE (CLIENTE):\n").bold = True
    p_cli.add_run("Cliente: {{ tercero.nom_tercero }}\nNIT/CC: {{ tercero.nit }}\nDirección: {{ tercero.direccion }}\nTeléfono: {{ tercero.telefono }}")

    doc.add_paragraph() # Espaciador

    # Tabla de Ítems (Con sintaxis Jinja {%tr for item in items %})
    table_items = doc.add_table(rows=2, cols=6)
    table_items.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ["#", "Ref.", "Descripción / Producto", "Cant.", "Vlr. Unitario", "Total (COP)"]
    hdr_cells = table_items.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        p = hdr_cells[i].paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        # Background color azul
        shd = docx.oxml.parse_xml(r'<w:shd {} w:fill="1D4ED8"/>'.format(docx.oxml.ns.nsdecls('w')))
        hdr_cells[i]._tc.get_or_add_tcPr().append(shd)

    # Fila plantilla con ciclo docxtpl {%tr for item in items %}
    row_cells = table_items.rows[1].cells
    row_cells[0].text = "{% for item in items %}{{ item.item_no }}"
    row_cells[1].text = "{{ item.referencia }}"
    row_cells[2].text = "{{ item.nom_articulo }}"
    row_cells[3].text = "{{ item.cantidad }}"
    row_cells[4].text = "{{ item.vlr_unitario }}"
    row_cells[5].text = "{{ item.tot_linea }}{% endfor %}"

    doc.add_paragraph() # Espaciador

    # Tabla de Totales y Código QR
    table_tot = doc.add_table(rows=1, cols=2)
    table_tot.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_qr = table_tot.rows[0].cells[0]
    cell_tot = table_tot.rows[0].cells[1]

    p_qr = cell_qr.paragraphs[0]
    p_qr.add_run("CÓDIGO QR DIAN:\n").bold = True
    p_qr.add_run("{{ qr_code }}\n\n")
    p_qr.add_run("FORMA / MEDIOS DE PAGO:\n").bold = True
    p_qr.add_run("{% for pago in pagos %}- {{ pago.medio }}: {{ pago.valor }}\n{% endfor %}")

    p_tot = cell_tot.paragraphs[0]
    p_tot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_tot.add_run("Subtotal Mercancía: {{ totals.tot_mercancia }}\n").font.size = Pt(10)
    p_tot.add_run("IVA Generado: {{ totals.tot_iva }}\n").font.size = Pt(10)
    p_tot.add_run("Retenciones (COP): {{ totals.tot_retefuente }}\n").font.size = Pt(10)
    r_gt = p_tot.add_run("TOTAL A PAGAR: {{ totals.tot_documento }}\n")
    r_gt.font.size = Pt(14)
    r_gt.font.bold = True
    r_gt.font.color.rgb = RGBColor(21, 128, 61) # Verde

    # Pie de página CUFE
    doc.add_paragraph()
    p_cufe = doc.add_paragraph()
    p_cufe.add_run("CUFE / Código Único de Factura Electrónica:\n").bold = True
    p_cufe.add_run("{{ doc.cufe }}").font.size = Pt(8)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Plantilla .docx creada exitosamente en: {output_path}")

if __name__ == '__main__':
    target = r"c:\Movil_apps\north_comercial\impresion\templates\impresion\plantilla_factura.docx"
    generate_invoice_template(target)
